# Ultralytics 🚀 AGPL-3.0 License - https://ultralytics.com/license

import os
import sys
import docx
import logging

# Ensure project root is in path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

DOCX_PATHS = [
    "/mnt/hdd/ultralytics/Computer_Vision_CCTV_Tracking_Project_Report.docx",
    "/mnt/hdd/ultralytics/project/Computer_Vision_CCTV_Tracking_Project_Report.docx"
]

def clean_details():
    for doc_path in DOCX_PATHS:
        if not os.path.exists(doc_path):
            logger.warning(f"File not found: {doc_path}")
            continue

        logger.info(f"Cleaning student details in: {doc_path}")
        doc = docx.Document(doc_path)
        
        # Clean paragraph 8: ERP ID
        p8 = doc.paragraphs[8]
        if "ERP ID: [241108757]" in p8.text:
            p8.text = "ERP ID: 241108757"
            logger.info("Cleaned ERP ID.")
            
        # Clean paragraph 9: PRN No
        p9 = doc.paragraphs[9]
        if "PRN No.: [2405020013588]" in p9.text:
            p9.text = "PRN No.: 2405020013588"
            logger.info("Cleaned PRN No.")

        doc.save(doc_path)
        logger.info(f"Saved cleaned details to: {doc_path}")

if __name__ == "__main__":
    clean_details()
