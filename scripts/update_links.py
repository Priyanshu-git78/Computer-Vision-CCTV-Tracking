# Ultralytics 🚀 AGPL-3.0 License - https://ultralytics.com/license

import os
import sys
import docx
import logging

# Ensure project root is in path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

DOCX_PATHS = [
    "/mnt/hdd/ultralytics/Computer_Vision_CCTV_Tracking_Project_Report.docx",
    "/mnt/hdd/ultralytics/project/Computer_Vision_CCTV_Tracking_Project_Report.docx"
]

def update_links():
    for doc_path in DOCX_PATHS:
        if not os.path.exists(doc_path):
            logger.warning(f"File not found: {doc_path}")
            continue

        logger.info(f"Updating dataset links and status in: {doc_path}")
        doc = docx.Document(doc_path)
        
        # 1. Update Table 10, Row 5 (Experience-center data)
        t10 = doc.tables[10]
        if len(t10.rows) >= 6:
            t10.cell(5, 1).text = "500 frames (simulated)"
            t10.cell(5, 2).text = "person, fire"
            t10.cell(5, 3).text = "YOLO txt"
            t10.cell(5, 4).text = "Internal / Authorized"
            t10.cell(5, 5).text = "Completed"
            logger.info("Updated Table 10 Row 5.")

        # 2. Update Table 18, Col 6 (Sources/Links)
        t18 = doc.tables[18]
        if len(t18.rows) >= 5:
            t18.cell(1, 6).text = "GitHub: github.com/Priyanshu-git78/Computer-Vision-CCTV-Tracking"
            t18.cell(2, 6).text = "GitHub: github.com/ultralytics/assets"
            t18.cell(3, 6).text = "Oxford: robots.ox.ac.uk/ActiveVision/Research/Projects/TownCentre"
            t18.cell(4, 6).text = "GitHub: github.com/Priyanshu-git78/Computer-Vision-CCTV-Tracking"
            logger.info("Updated Table 18 Source links.")

        doc.save(doc_path)
        logger.info(f"Saved updated document to: {doc_path}")

if __name__ == "__main__":
    update_links()
