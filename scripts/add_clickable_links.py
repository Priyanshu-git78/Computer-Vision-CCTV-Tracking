# Ultralytics 🚀 AGPL-3.0 License - https://ultralytics.com/license

import os
import sys
import docx
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Ensure project root is in path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

DOCX_PATHS = [
    "/mnt/hdd/ultralytics/Computer_Vision_CCTV_Tracking_Project_Report.docx",
    "/mnt/hdd/ultralytics/project/Computer_Vision_CCTV_Tracking_Project_Report.docx"
]

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function that places a clickable hyperlink inside a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create w:hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create w:r
    new_run = OxmlElement('w:r')

    # Create w:rPr
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def format_docx_links():
    for doc_path in DOCX_PATHS:
        if not os.path.exists(doc_path):
            logger.warning(f"File not found: {doc_path}")
            continue

        logger.info(f"Adding clickable hyperlinks to: {doc_path}")
        doc = docx.Document(doc_path)

        # 1. Format Paragraph 447
        p447 = doc.paragraphs[447]
        p447.text = "" # Clear paragraph
        
        run1 = p447.add_run("Repository Link: ")
        run1.bold = True
        add_hyperlink(p447, "https://github.com/Priyanshu-git78/Computer-Vision-CCTV-Tracking", "https://github.com/Priyanshu-git78/Computer-Vision-CCTV-Tracking")
        
        p447.add_run("\n") # New line
        
        run2 = p447.add_run("Hugging Face Model Hub: ")
        run2.bold = True
        add_hyperlink(p447, "https://huggingface.co/Priyanshu-68/yolov8-cctv-tracking-models", "https://huggingface.co/Priyanshu-68/yolov8-cctv-tracking-models")
        logger.info("Formatted Paragraph 447 clickable links.")

        # 2. Format Table 18 Source Links
        t18 = doc.tables[18]
        links_map = {
            1: ("https://github.com/Priyanshu-git78/Computer-Vision-CCTV-Tracking", "GitHub: Computer-Vision-CCTV-Tracking"),
            2: ("https://github.com/ultralytics/assets", "GitHub: ultralytics/assets"),
            3: ("https://www.robots.ox.ac.uk/ActiveVision/Research/Projects/TownCentre/index.xml", "Oxford: TownCentre Dataset"),
            4: ("https://github.com/Priyanshu-git78/Computer-Vision-CCTV-Tracking", "GitHub: Computer-Vision-CCTV-Tracking")
        }

        for row_idx, (url, text) in links_map.items():
            cell = t18.cell(row_idx, 6)
            cell.text = "" # Clear cell text
            p = cell.paragraphs[0]
            add_hyperlink(p, url, text)
            
        logger.info("Formatted Table 18 clickable links.")

        doc.save(doc_path)
        logger.info(f"Successfully saved: {doc_path}")

if __name__ == "__main__":
    format_docx_links()
